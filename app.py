import os
import streamlit as st
from google import generativeai as genai
import pdfplumber
import docx

# --- Language Translations ---
translations = {
    ":gb: English": {
        "page_title": "Job Application Booster",
        "page_icon": "💼",
        "title": "💼 Job Application Booster",
        "subheader": "Your AI Assistant for Tailored Applications",
        "settings_header": "⚙️ Settings",
        "language_label": "🌍 Language",
        "tone_label": "🎭 Output Tone",
        "tones": ["Professional", "Friendly", "Persuasive", "Confident"],
        "tone_prompts": {
            "Professional": "Ensure your tone is formal, respectful, and industry-appropriate.",
            "Friendly": "Make the response sound warm, approachable, and conversational.",
            "Persuasive": "Use convincing, enthusiastic language that sells the candidate strongly.",
            "Confident": "Make the candidate sound self-assured, competent, and assertive."
        },
        "creativity_label": "🎛️ Model Creativity",
        "sections_header": "🧩 Choose Sections to Generate",
        "sections": {
            "cl": "✉️ Cover Letter",
            "ri": "📌 Resume Improvements",
            "it": "💡 Interview Tips",
            "cv": "🔄 Rewrite My CV"
        },
        "upload_header": "1. Upload Your Files",
        "resume_upload": "📄 Upload Resume (.pdf or .docx)",
        "jd_upload": "📌 Upload Job Description (.pdf or .docx)",
        "resume_paste": "✍️ Paste your Resume (or review extracted)",
        "jd_paste": "📋 Paste Job Description (or review extracted)",
        "generate_header": "2. Boost!",
        "generate_button": "🚀 Boost My Application",
        "custom_prompt_header": "🛠️ Option 2: Adjust Cover Letter Output",
        "custom_prompt_placeholder": "e.g. Make it more persuasive and highlight leadership achievements",
        "error_api": "❌ API key not found. Please configure your secrets.toml file like:\n\n[google]\napi_key = \"your_real_api_key_here\"",
        "error_config": "❌ Gemini API Configuration Failed: {}",
        "error_resume": "❌ Error reading resume: {}",
        "error_jd": "❌ Error reading job description: {}",
        "warning_input": "⚠️ Please provide both Resume and Job Description.",
        "warning_sections": "⚠️ Please select at least one generation option.",
        "warning_custom": "⚠️ Resume and Job Description must be filled.",
        "regenerating": "🔁 Regenerating Cover Letter...",
        "generating": "✨ Generating your application materials...",
        "error_generation": "⚠️ Error regenerating Cover Letter: {}",
        "error_api_generation": "❌ Gemini API Error: {}",
        "footer": "Built with ❤️ by Okoli Chidozie (GoldenBlaze)"
    },
    ":es: Spanish": {
        "page_title": "Potenciador de Solicitudes de Empleo",
        "page_icon": "💼",
        "title": "💼 Potenciador de Solicitudes de Empleo",
        "subheader": "Tu Asistente de IA para Aplicaciones Personalizadas",
        "settings_header": "⚙️ Configuración",
        "language_label": "🌍 Idioma",
        "tone_label": "🎭 Tono de Salida",
        "tones": ["Profesional", "Amigable", "Persuasivo", "Seguro"],
        "tone_prompts": {
            "Profesional": "Asegúrate de que el tono sea formal, respetado y adecuado para la industria.",
            "Amigable": "Haz que la respuesta suene cálida, accesible y conversacional.",
            "Persuasivo": "Usa un lenguaje convincente y entusiasta que venda fuertemente al candidato.",
            "Seguro": "Haz que el candidato suene seguro de sí mismo, competente y asertivo."
        },
        "creativity_label": "🎛️ Creatividad del Modelo",
        "sections_header": "🧩 Elige Secciones a Generar",
        "sections": {
            "cl": "✉️ Carta de Presentación",
            "ri": "📌 Mejoras para el CV",
            "it": "💡 Consejos para Entrevistas",
            "cv": "🔄 Reescribir mi CV"
        },
        "upload_header": "1. Sube Tus Archivos",
        "resume_upload": "📄 Subir CV (.pdf o .docx)",
        "jd_upload": "📌 Subir Descripción del Puesto (.pdf o .docx)",
        "resume_paste": "✍️ Pega tu CV (o revisa el extraído)",
        "jd_paste": "📋 Pega la Descripción del Puesto (o revisa la extraída)",
        "generate_header": "2. ¡Potencia!",
        "generate_button": "🚀 Potenciar mi Aplicación",
        "custom_prompt_header": "🛠️ Opción 2: Ajustar la Carta de Presentación",
        "custom_prompt_placeholder": "ej. Hazla más persuasiva y destaca logros de liderazgo",
        "error_api": "❌ Clave API no encontrada. Configura tu archivo secrets.toml así:\n\n[google]\napi_key = \"tu_clave_api_real_aqui\"",
        "error_config": "❌ Configuración de API Gemini fallida: {}",
        "error_resume": "❌ Error leyendo CV: {}",
        "error_jd": "❌ Error leyendo descripción del puesto: {}",
        "warning_input": "⚠️ Proporciona tanto tu CV como la Descripción del Puesto.",
        "warning_sections": "⚠️ Selecciona al menos una opción de generación.",
        "warning_custom": "⚠️ Debes completar tanto el CV como la Descripción del Puesto.",
        "regenerating": "🔁 Regenerando Carta de Presentación...",
        "generating": "✨ Generando tus materiales de aplicación...",
        "error_generation": "⚠️ Error regenerando Carta de Presentación: {}",
        "error_api_generation": "❌ Error de API Gemini: {}",
        "footer": "Desarrollado con ❤️ por Okoli Chidozie (GoldenBlaze)"
    },
    ":fr: French": {
        "page_title": "Booster de Candidature",
        "page_icon": "💼",
        "title": "💼 Booster de Candidature",
        "subheader": "Votre Assistant IA pour des Candidatures Sur Mesure",
        "settings_header": "⚙️ Paramètres",
        "language_label": "🌍 Langue",
        "tone_label": "🎭 Ton de Sortie",
        "tones": ["Professionnel", "Amical", "Persuasif", "Confiant"],
        "tone_prompts": {
            "Professionnel": "Assurez-vous que le ton est formel, respectueux et adapté à l'industrie.",
            "Amical": "Rendez la réponse chaleureuse, abordable et conversationnelle.",
            "Persuasif": "Utilisez un langage convaincant et enthousiaste qui met fortement en valeur le candidat.",
            "Confiant": "Faites en sorte que le candidat semble sûr de lui, compétent et assertif."
        },
        "creativity_label": "🎛️ Créativité du Modèle",
        "sections_header": "🧩 Choisir les Sections à Générer",
        "sections": {
            "cl": "✉️ Lettre de Motivation",
            "ri": "📌 Améliorations de CV",
            "it": "💡 Conseils d'Entretien",
            "cv": "🔄 Réécrire mon CV"
        },
        "upload_header": "1. Téléchargez Vos Fichiers",
        "resume_upload": "📄 Télécharger CV (.pdf ou .docx)",
        "jd_upload": "📌 Télécharger Description de Poste (.pdf ou .docx)",
        "resume_paste": "✍️ Collez votre CV (ou vérifiez l'extrait)",
        "jd_paste": "📋 Collez la Description de Poste (ou vérifiez l'extrait)",
        "generate_header": "2. Boostez!",
        "generate_button": "🚀 Booster ma Candidature",
        "custom_prompt_header": "🛠️ Option 2: Ajuster la Lettre de Motivation",
        "custom_prompt_placeholder": "ex. Rendez-la plus persuasive et mettez en avant les réalisations en leadership",
        "error_api": "❌ Clé API introuvable. Configurez votre fichier secrets.toml comme suit:\n\n[google]\napi_key = \"votre_vraie_clé_api_ici\"",
        "error_config": "❌ Échec de Configuration de l'API Gemini: {}",
        "error_resume": "❌ Erreur de lecture du CV: {}",
        "error_jd": "❌ Erreur de lecture de la description de poste: {}",
        "warning_input": "⚠️ Veuillez fournir à la fois votre CV et la Description de Poste.",
        "warning_sections": "⚠️ Veuillez sélectionner au moins une option de génération.",
        "warning_custom": "⚠️ Le CV et la Description de Poste doivent être remplis.",
        "regenerating": "🔁 Régénération de la Lettre de Motivation...",
        "generating": "✨ Génération de vos documents de candidature...",
        "error_generation": "⚠️ Erreur lors de la régénération de la Lettre de Motivation: {}",
        "error_api_generation": "❌ Erreur de l'API Gemini: {}",
        "footer": "Construit avec ❤️ par Okoli Chidozie (GoldenBlaze)"
    },
    ":de: German": {
        "page_title": "Job Bewerbungs-Booster",
        "page_icon": "💼",
        "title": "💼 Job Bewerbungs-Booster",
        "subheader": "Ihr KI-Assistent für maßgeschneiderte Bewerbungen",
        "settings_header": "⚙️ Einstellungen",
        "language_label": "🌍 Sprache",
        "tone_label": "🎭 Ausgabeton",
        "tones": ["Professionell", "Freundlich", "Überzeugend", "Selbstbewusst"],
        "tone_prompts": {
            "Professionell": "Stellen Sie sicher, dass der Ton formal, respektvoll und branchengerecht ist.",
            "Freundlich": "Lassen Sie die Antwort warm, zugänglich und gesprächsweise klingen.",
            "Überzeugend": "Verwenden Sie überzeugende, begeisternde Sprache, die den Kandidaten stark verkauft.",
            "Selbstbewusst": "Lassen Sie den Kandidaten selbstsicher, kompetent und durchsetzungsfähig klingen."
        },
        "creativity_label": "🎛️ Modellkreativität",
        "sections_header": "🧩 Wählen Sie zu generierende Abschnitte",
        "sections": {
            "cl": "✉️ Anschreiben",
            "ri": "📌 Lebenslauf-Verbesserungen",
            "it": "💡 Vorstellungsgespräch-Tipps",
            "cv": "🔄 Lebenslauf umschreiben"
        },
        "upload_header": "1. Laden Sie Ihre Dateien hoch",
        "resume_upload": "📄 Lebenslauf hochladen (.pdf oder .docx)",
        "jd_upload": "📌 Stellenbeschreibung hochladen (.pdf oder .docx)",
        "resume_paste": "✍️ Fügen Sie Ihren Lebenslauf ein (oder überprüfen Sie die extrahierte Version)",
        "jd_paste": "📋 Fügen Sie die Stellenbeschreibung ein (oder überprüfen Sie die extrahierte Version)",
        "generate_header": "2. Boost!",
        "generate_button": "🚀 Meine Bewerbung boosten",
        "custom_prompt_header": "🛠️ Option 2: Anschreiben anpassen",
        "custom_prompt_placeholder": "z.B. Machen Sie es überzeugender und heben Sie Führungsleistungen hervor",
        "error_api": "❌ API-Schlüssel nicht gefunden. Bitte konfigurieren Sie Ihre secrets.toml-Datei wie folgt:\n\n[google]\napi_key = \"Ihr_echter_api_schluessel_hier\"",
        "error_config": "❌ Gemini API Konfiguration fehlgeschlagen: {}",
        "error_resume": "❌ Fehler beim Lesen des Lebenslaufs: {}",
        "error_jd": "❌ Fehler beim Lesen der Stellenbeschreibung: {}",
        "warning_input": "⚠️ Bitte geben Sie sowohl Lebenslauf als auch Stellenbeschreibung an.",
        "warning_sections": "⚠️ Bitte wählen Sie mindestens eine Generierungsoption.",
        "warning_custom": "⚠️ Lebenslauf und Stellenbeschreibung müssen ausgefüllt sein.",
        "regenerating": "🔁 Anschreiben wird neu generiert...",
        "generating": "✨ Ihre Bewerbungsunterlagen werden generiert...",
        "error_generation": "⚠️ Fehler beim Neugenerieren des Anschreibens: {}",
        "error_api_generation": "❌ Gemini API Fehler: {}",
        "footer": "Erstellt mit ❤️ von Okoli Chidozie (GoldenBlaze)"
    }
}

# --- API Key from secrets.toml ---
try:
    GOOGLE_API_KEY = st.secrets["google"]["api_key"]
    genai.configure(api_key=GOOGLE_API_KEY)
    model = genai.GenerativeModel("gemini-1.5-flash")
except KeyError:
    st.error(translations[":gb: English"]["error_api"])
    st.stop()
except Exception as e:
    st.error(translations[":gb: English"]["error_config"].format(e))
    st.stop()

# === PAGE CONFIGURATION (must be first Streamlit command) ===
default_language = ":gb: English"
trans = translations[default_language]
st.set_page_config(
    page_title=trans["page_title"],
    page_icon=trans["page_icon"]
)

# === SIDEBAR CONFIGURATION ===
st.sidebar.header(trans["settings_header"])

# Language Selector (at the top of sidebar)
language = st.sidebar.selectbox(
    trans["language_label"],
    list(translations.keys()),
    key="language",
    format_func=lambda x: x.replace(":gb:", "🇬🇧")
                         .replace(":es:", "🇪🇸")
                         .replace(":fr:", "🇫🇷")
                         .replace(":de:", "🇩🇪")
)

# Update translations without recalling set_page_config
trans = translations[language]

# Tone Selector
tone = st.sidebar.selectbox(
    trans["tone_label"],
    trans["tones"],
    key="tone"
)
tone_instruction = trans["tone_prompts"][tone]

# Creativity Level
st.sidebar.slider(
    trans["creativity_label"],
    0.0, 1.0, 0.7, 0.1,
    key="temperature"
)

# Section Selection
st.sidebar.header(trans["sections_header"])
generate_cl = st.sidebar.checkbox(trans["sections"]["cl"], value=True)
generate_ri = st.sidebar.checkbox(trans["sections"]["ri"], value=True)
generate_it = st.sidebar.checkbox(trans["sections"]["it"], value=True)
generate_cv = st.sidebar.checkbox(trans["sections"]["cv"], value=False)

# === MAIN AREA ===
st.title(trans["title"])
st.subheader(trans["subheader"])

# --- 1. Upload Resume & Job Description ---
st.header(trans["upload_header"])
col1, col2 = st.columns(2)

with col1:
    uploaded_resume = st.file_uploader(
        trans["resume_upload"],
        type=["pdf", "docx"]
    )
with col2:
    uploaded_jd = st.file_uploader(
        trans["jd_upload"],
        type=["pdf", "docx"]
    )

# --- Extract Resume Text ---
extracted_resume = ""
if uploaded_resume:
    try:
        if uploaded_resume.name.endswith("pdf"):
            with pdfplumber.open(uploaded_resume) as pdf:
                extracted_resume = "\n".join([page.extract_text() or "" for page in pdf.pages])
        elif uploaded_resume.name.endswith("docx"):
            doc = docx.Document(uploaded_resume)
            extracted_resume = "\n".join([para.text for para in doc.paragraphs])
    except Exception as e:
        st.error(trans["error_resume"].format(e))

# --- Extract JD Text ---
extracted_jd = ""
if uploaded_jd:
    try:
        if uploaded_jd.name.endswith("pdf"):
            with pdfplumber.open(uploaded_jd) as pdf:
                extracted_jd = "\n".join([page.extract_text() or "" for page in pdf.pages])
        elif uploaded_jd.name.endswith("docx"):
            doc = docx.Document(uploaded_jd)
            extracted_jd = "\n".join([para.text for para in doc.paragraphs])
    except Exception as e:
        st.error(trans["error_jd"].format(e))

# --- Paste & Edit ---
resume = st.text_area(
    trans["resume_paste"],
    value=extracted_resume,
    height=250
)
job_desc = st.text_area(
    trans["jd_paste"],
    value=extracted_jd,
    height=250
)

# --- Generate Button ---
st.header(trans["generate_header"])
start = st.button(trans["generate_button"])

# --- Prompt box to edit Cover Letter specifically ---
if generate_cl:
    st.markdown(f"### {trans['custom_prompt_header']}")
    user_custom_prompt = st.text_input(
        trans["custom_prompt_placeholder"],
        key="custom_prompt"
    )
    if user_custom_prompt:
        if not resume or not job_desc:
            st.warning(trans["warning_custom"])
        else:
            # Extract just the language name (without flag)
            language_name = language.split(":")[-1].strip()
            prompt = f"""
            You are a job application writing expert. Respond in {language_name}.

            Resume:
            {resume}

            Job Description:
            {job_desc}

            Tone: {tone} — {tone_instruction}

            Now regenerate the Cover Letter using this instruction:
            "{user_custom_prompt}"

            Output only the ### Cover Letter section in {language_name}.
            """
            with st.spinner(trans["regenerating"]):
                try:
                    stream = model.generate_content(
                        prompt,
                        stream=True,
                        generation_config={"temperature": st.session_state.temperature}
                    )
                    cl_out = ""
                    cl_box = st.empty()
                    for chunk in stream:
                        cl_out += chunk.text
                        cl_box.markdown(cl_out + "▌")
                except Exception as e:
                    st.error(trans["error_generation"].format(e))

# --- Main Generation ---
if start and not user_custom_prompt:
    if not resume or not job_desc:
        st.warning(trans["warning_input"])
    elif not (generate_cl or generate_ri or generate_it or generate_cv):
        st.warning(trans["warning_sections"])
    else:
        # Extract just the language name (without flag)
        language_name = language.split(":")[-1].strip()
        with st.spinner(trans["generating"]):
            parts = []
            if generate_cl:
                parts.append(f"### {trans['sections']['cl']}\nWrite a tailored cover letter to 'Hiring Manager' in {language_name}.")
            if generate_ri:
                parts.append(f"### {trans['sections']['ri']}\nList 3–5 actionable suggestions based on the job description in {language_name}.")
            if generate_it:
                parts.append(f"### {trans['sections']['it']}\nList 3–5 concise, role-specific interview tips in {language_name}.")
            if generate_cv:
                parts.append(f"### {trans['sections']['cv']}\nGenerate a full rewritten version of the resume tailored to the job description using professional formatting in {language_name}.")

            full_prompt = f"""
            You are a professional career assistant helping users apply for jobs. Respond in {language_name}.

            Resume:
            {resume}

            Job Description:
            {job_desc}

            Tone: {tone} — {tone_instruction}

            Now generate the following sections using markdown ### headings in {language_name}:
            {" ".join(parts)}
            """

            try:
                stream = model.generate_content(
                    full_prompt,
                    stream=True,
                    generation_config={"temperature": st.session_state.temperature}
                )
                output = ""
                response_area = st.empty()
                for chunk in stream:
                    output += chunk.text
                    response_area.markdown(output + "▌")
            except Exception as e:
                st.error(trans["error_api_generation"].format(e))

# --- Footer ---
st.markdown("---")
st.caption(trans["footer"])